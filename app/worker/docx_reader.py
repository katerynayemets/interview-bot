from __future__ import annotations

from io import BytesIO
from docx import Document


def read_docx_text(file_bytes: bytes) -> str:
    """Extract plain text from a .docx file.

    Notes:
    - Supports paragraphs and tables.
    - No OCR (images/scans inside DOCX will return empty).
    """
    bio = BytesIO(file_bytes)
    doc = Document(bio)

    parts: list[str] = []

    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t:
            parts.append(t)

    for table in doc.tables:
        for row in table.rows:
            row_cells = []
            for cell in row.cells:
                t = (cell.text or "").strip()
                if t:
                    row_cells.append(t)
            if row_cells:
                parts.append(" | ".join(row_cells))

    return "\n".join(parts).strip()
